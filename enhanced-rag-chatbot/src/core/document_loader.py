import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
import streamlit as st
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader, CSVLoader, UnstructuredWordDocumentLoader
)
from PIL import Image
import numpy as np

# Make OpenCV import optional for headless environments
try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False
    print("OpenCV not available. Image analysis features will be limited.")

class EnhancedDocumentLoader:
    """Enhanced document loader with multi-modal support"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._load_pdf,
            'txt': self._load_text,
            'csv': self._load_csv,
            'docx': self._load_docx,
            'png': self._load_image,
            'jpg': self._load_image,
            'jpeg': self._load_image
        }
    
    def load_documents(self, uploaded_files: List) -> List[Document]:
        """Load documents from uploaded files"""
        documents = []
        
        for uploaded_file in uploaded_files:
            try:
                # Create temporary file
                with tempfile.NamedTemporaryFile(
                    delete=False,
                    suffix=f".{uploaded_file.name.split('.')[-1]}"
                ) as tmp_file:
                    tmp_file.write(uploaded_file.read())
                    tmp_path = tmp_file.name
                
                # Get file extension
                extension = uploaded_file.name.split('.')[-1].lower()
                
                if extension in self.supported_formats:
                    docs = self.supported_formats[extension](tmp_path, uploaded_file.name)
                    documents.extend(docs)
                else:
                    st.warning(f"Unsupported file format: {extension}")
                
                # Clean up
                Path(tmp_path).unlink(missing_ok=True)
                
            except Exception as e:
                st.error(f"Error loading {uploaded_file.name}: {str(e)}")
        
        return documents
    
    def _load_pdf(self, file_path: str, filename: str) -> List[Document]:
        """Load PDF documents"""
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            for doc in documents:
                doc.metadata.update({
                    "source": filename,
                    "file_type": "pdf"
                })
            
            return documents
        except Exception as e:
            st.error(f"Error loading PDF {filename}: {str(e)}")
            return []
    
    def _load_text(self, file_path: str, filename: str) -> List[Document]:
        """Load text files"""
        try:
            try:
                loader = TextLoader(file_path, encoding='utf-8')
                documents = loader.load()
            except UnicodeDecodeError:
                # Try with different encoding
                loader = TextLoader(file_path, encoding='latin-1')
                documents = loader.load()
            
            for doc in documents:
                doc.metadata.update({
                    "source": filename,
                    "file_type": "text"
                })
            
            return documents
        except Exception as e:
            st.error(f"Error loading text file {filename}: {str(e)}")
            return []
    
    def _load_csv(self, file_path: str, filename: str) -> List[Document]:
        """Load CSV files"""
        try:
            loader = CSVLoader(file_path)
            documents = loader.load()
            
            for doc in documents:
                doc.metadata.update({
                    "source": filename,
                    "file_type": "csv"
                })
            
            return documents
        except Exception as e:
            st.error(f"Error loading CSV {filename}: {str(e)}")
            return []
    
    def _load_docx(self, file_path: str, filename: str) -> List[Document]:
        """Load Word documents"""
        try:
            try:
                loader = UnstructuredWordDocumentLoader(file_path)
                documents = loader.load()
            except Exception:
                # Fallback to simple text extraction
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    content = "\n".join([para.text for para in doc.paragraphs])
                    documents = [Document(page_content=content)]
                except Exception:
                    # Final fallback - create empty document
                    content = f"Could not extract text from {filename}"
                    documents = [Document(page_content=content)]
            
            for doc in documents:
                doc.metadata.update({
                    "source": filename,
                    "file_type": "docx"
                })
            
            return documents
        except Exception as e:
            st.error(f"Error loading DOCX {filename}: {str(e)}")
            return []
    
    def _load_image(self, file_path: str, filename: str) -> List[Document]:
        """Load and process images"""
        try:
            # Load image with PIL
            image = Image.open(file_path)
            
            if CV2_AVAILABLE:
                # Use OpenCV for advanced processing
                img_array = np.array(image)
                height, width = img_array.shape[:2]
            else:
                # Fallback to PIL
                width, height = image.size
            
            # Create basic document
            content = f"Image: {filename}\nDimensions: {width}x{height}\nFormat: {image.format}"
            
            # Add basic image analysis
            if hasattr(image, 'mode'):
                content += f"\nColor Mode: {image.mode}"
            
            document = Document(
                page_content=content,
                metadata={
                    "source": filename,
                    "file_type": "image",
                    "width": width,
                    "height": height,
                    "format": image.format
                }
            )
            
            return [document]
            
        except Exception as e:
            st.error(f"Error processing image {filename}: {str(e)}")
            return []
